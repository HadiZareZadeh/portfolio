# -*- coding: utf-8 -*-
"""
Generate a Word (.docx) resume matching the exact style specification.
Run: python generate_resume_docx.py
Output: resume_data_engineer.docx
"""

from docx import Document
from docx.shared import Cm, Pt, Inches, Twips, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Colors (hex without # for OOXML)
DARK_BLUE = "1F4E5F"
BROWN = "A66A2C"
BODY_TEXT = "4A4A4A"
LIGHT_GRAY = "8A8A8A"
DIVIDER = "D8CFC2"
BACKGROUND = "F6F1E6"

FONT_NAME = "Georgia"


def set_document_background(doc, hex_color):
    """Set page background color via OOXML."""
    # w:background must be first child of w:body
    body = doc.element.body
    bg = OxmlElement("w:background")
    bg.set(qn("w:color"), hex_color)
    body.insert(0, bg)
    # Enable display of background in Word
    try:
        settings = doc.settings.element
        disp = OxmlElement("w:displayBackgroundShape")
        settings.insert(0, disp)
    except Exception:
        pass


def rgb_from_hex(hex_color):
    """Convert hex (without #) to (r, g, b) 0-255."""
    h = hex_color.lstrip("#")
    return tuple(int(h[i : i + 2], 16) for i in (0, 2, 4))


def add_paragraph_with_style(doc, text, font_size_pt, bold=False, color_hex=None, font_name=FONT_NAME):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(font_size_pt)
    run.font.bold = bold
    run.font.name = font_name
    if color_hex:
        r, g, b = rgb_from_hex(color_hex)
        run.font.color.rgb = RGBColor(r, g, b)
    return p


def add_divider(doc, space_before_pt=8, space_after_pt=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before_pt)
    p.paragraph_format.space_after = Pt(space_after_pt)
    p.paragraph_format.border_bottom_width = Pt(0.75)
    p.paragraph_format.border_bottom_color = rgb_from_hex(DIVIDER)
    # In python-docx, border color is set via XML
    p_border = p._p.get_or_add_pPr()
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")  # 0.75 pt approx
    bottom.set(qn("w:color"), DIVIDER)
    p_border.append(bottom)
    return p


def set_cell_shading(cell, hex_color):
    """Set cell shading (for possible use)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def run_style(run, size_pt, bold=False, color_hex=None):
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.name = FONT_NAME
    if color_hex:
        r, g, b = rgb_from_hex(color_hex)
        run.font.color.rgb = RGBColor(r, g, b)


def main():
    doc = Document()
    doc.styles["Normal"].font.name = FONT_NAME
    doc.styles["List Paragraph"].font.name = FONT_NAME

    # ----- Page setup -----
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    set_document_background(doc, BACKGROUND)

    # ----- Header (full width above columns) -----
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(10)
    r = title.add_run("Data Engineer")
    run_style(r, 22, bold=True, color_hex=DARK_BLUE)

    # Divider line under title (full width, 0.75 pt, divider color)
    div = doc.add_paragraph()
    div.paragraph_format.space_before = Pt(8)
    div.paragraph_format.space_after = Pt(8)
    pPr = div._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    b = OxmlElement("w:bottom")
    b.set(qn("w:val"), "single")
    b.set(qn("w:sz"), "6")
    b.set(qn("w:color"), DIVIDER)
    pBdr.append(b)
    pPr.append(pBdr)

    # ----- Two-column layout via table -----
    # One row, two cells. Left 62%, Right 38%. No borders.
    total_width_cm = 21.0 - 2.0 - 2.0  # page - left margin - right margin
    left_width_cm = total_width_cm * 0.62
    right_width_cm = total_width_cm * 0.38
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.allow_autofit = False
    # Column widths in twips (1 cm ≈ 567 twips)
    table.columns[0].width = int(left_width_cm * 567)
    table.columns[1].width = int(right_width_cm * 567)
    def make_nil_borders():
        borders_el = OxmlElement("w:tcBorders")
        for side in ["top", "left", "bottom", "right"]:
            b_el = OxmlElement(f"w:{side}")
            b_el.set(qn("w:val"), "nil")
            borders_el.append(b_el)
        return borders_el
    for row in table.rows:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcPr.append(make_nil_borders())
    left_cell = table.rows[0].cells[0]
    right_cell = table.rows[0].cells[1]

    # ----- LEFT COLUMN -----
    left_doc = left_cell

    # Contact block (icons + text, 10.5 pt, 6 pt between lines)
    contact_items = [
        ("●", "Martha Stewart"),           # Person
        ("☎", "555-555-5555"),             # Phone
        ("◆", "Vancouver, British Columbia"),  # Location
        ("✉", "email@example.com"),       # Email
        ("◉", "example.com"),              # Website
    ]
    for icon_char, text in contact_items:
        p = left_doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = Pt(10.5 + 2)
        ri = p.add_run(icon_char + "  ")
        run_style(ri, 10.5, color_hex=BROWN)
        rt = p.add_run(text)
        run_style(rt, 10.5, color_hex=BODY_TEXT)

    # Spacing after contact
    spacer = left_doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(12)

    # Section: Work Experience
    p_we = left_doc.add_paragraph()
    p_we.paragraph_format.space_before = Pt(4)
    r_we = p_we.add_run("Work Experience")
    run_style(r_we, 15, bold=True, color_hex=DARK_BLUE)
    p_we.paragraph_format.space_after = Pt(4)
    # Divider under section title
    div_we = left_doc.add_paragraph()
    div_we.paragraph_format.space_before = Pt(0)
    div_we.paragraph_format.space_after = Pt(8)
    pPr = div_we._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    b = OxmlElement("w:bottom")
    b.set(qn("w:val"), "single")
    b.set(qn("w:sz"), "6")
    b.set(qn("w:color"), DIVIDER)
    pBdr.append(b)
    pPr.append(pBdr)

    # Job 1
    p_d1 = left_doc.add_paragraph()
    r_d1 = p_d1.add_run("2018-10 – Present")
    run_style(r_d1, 9, color_hex=LIGHT_GRAY)
    p_d1.paragraph_format.space_after = Pt(2)
    p_t1 = left_doc.add_paragraph()
    r_t1 = p_t1.add_run("Architect/Data Engineer")
    run_style(r_t1, 12, bold=True, color_hex=DARK_BLUE)
    p_t1.paragraph_format.space_after = Pt(0)
    p_c1 = left_doc.add_paragraph()
    r_c1 = p_c1.add_run("Amazon.com")
    run_style(r_c1, 10.5, color_hex=BROWN)
    p_c1.paragraph_format.space_after = Pt(4)
    bullets1 = [
        "MapReduce jobs in Java & Spark",
        "Kafka pipelines",
        "HBase, Elasticsearch",
        "ETL processes",
        "Python frameworks",
        "Real-time & batch processing",
        "Business insights & analytics impact",
    ]
    for bullet in bullets1:
        bp = left_doc.add_paragraph(style="List Bullet")
        bp.paragraph_format.space_after = Pt(2)
        bp.paragraph_format.line_spacing = 1.25
        bp.paragraph_format.left_indent = Cm(0.5)
        r = bp.add_run(bullet)
        run_style(r, 10.5, color_hex=BODY_TEXT)

    # Job 2
    p_d2 = left_doc.add_paragraph()
    p_d2.paragraph_format.space_before = Pt(12)
    r_d2 = p_d2.add_run("2008-02 – 2018-09")
    run_style(r_d2, 9, color_hex=LIGHT_GRAY)
    p_d2.paragraph_format.space_after = Pt(2)
    p_t2 = left_doc.add_paragraph()
    r_t2 = p_t2.add_run("Data Engineer")
    run_style(r_t2, 12, bold=True, color_hex=DARK_BLUE)
    p_t2.paragraph_format.space_after = Pt(0)
    p_c2 = left_doc.add_paragraph()
    r_c2 = p_c2.add_run("Meta")
    run_style(r_c2, 10.5, color_hex=BROWN)
    p_c2.paragraph_format.space_after = Pt(4)
    bullets2 = [
        "Large-scale data pipelines and warehousing",
        "Collaboration with analytics and product teams",
        "Performance tuning and optimization",
    ]
    for bullet in bullets2:
        bp = left_doc.add_paragraph(style="List Bullet")
        bp.paragraph_format.space_after = Pt(2)
        bp.paragraph_format.line_spacing = 1.25
        bp.paragraph_format.left_indent = Cm(0.5)
        r = bp.add_run(bullet)
        run_style(r, 10.5, color_hex=BODY_TEXT)

    # ----- RIGHT COLUMN -----
    right_doc = right_cell

    # ETL Portfolio
    p_etl = right_doc.add_paragraph()
    p_etl.paragraph_format.space_before = Pt(0)
    r_etl = p_etl.add_run("ETL Portfolio")
    run_style(r_etl, 15, bold=True, color_hex=DARK_BLUE)
    p_etl.paragraph_format.space_after = Pt(4)
    div_etl = right_doc.add_paragraph()
    div_etl.paragraph_format.space_after = Pt(8)
    pPr = div_etl._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    b = OxmlElement("w:bottom")
    b.set(qn("w:val"), "single")
    b.set(qn("w:sz"), "6")
    b.set(qn("w:color"), DIVIDER)
    pBdr.append(b)
    pPr.append(pBdr)
    p_etl_content = right_doc.add_paragraph()
    p_etl_content.paragraph_format.space_after = Pt(8)
    p_etl_content.paragraph_format.line_spacing = 1.25
    r_etl_pre = p_etl_content.add_run("ETL for Azure data lake: ")
    run_style(r_etl_pre, 10.5, color_hex=BODY_TEXT)
    r_etl_link = p_etl_content.add_run("www.example.com/etl")
    run_style(r_etl_link, 10.5, color_hex=BROWN)

    # Summary
    p_sum = right_doc.add_paragraph()
    p_sum.paragraph_format.space_before = Pt(4)
    r_sum = p_sum.add_run("Summary")
    run_style(r_sum, 15, bold=True, color_hex=DARK_BLUE)
    p_sum.paragraph_format.space_after = Pt(4)
    div_sum = right_doc.add_paragraph()
    div_sum.paragraph_format.space_after = Pt(8)
    pPr = div_sum._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    b = OxmlElement("w:bottom")
    b.set(qn("w:val"), "single")
    b.set(qn("w:sz"), "6")
    b.set(qn("w:color"), DIVIDER)
    pBdr.append(b)
    pPr.append(pBdr)
    summary_text = (
        "Data engineer with extensive experience building and maintaining large-scale data pipelines, "
        "ETL processes, and analytics infrastructure. Skilled in MapReduce, Spark, Kafka, and cloud data lakes. "
        "Focused on delivering reliable, scalable solutions that drive business insights and decision-making."
    )
    p_sum_p = right_doc.add_paragraph()
    p_sum_p.paragraph_format.space_after = Pt(8)
    p_sum_p.paragraph_format.line_spacing = 1.3
    p_sum_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r_sum_p = p_sum_p.add_run(summary_text)
    run_style(r_sum_p, 10.5, color_hex=BODY_TEXT)

    # Education
    p_edu = right_doc.add_paragraph()
    p_edu.paragraph_format.space_before = Pt(4)
    r_edu = p_edu.add_run("Education")
    run_style(r_edu, 15, bold=True, color_hex=DARK_BLUE)
    p_edu.paragraph_format.space_after = Pt(4)
    div_edu = right_doc.add_paragraph()
    div_edu.paragraph_format.space_after = Pt(8)
    pPr = div_edu._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    b = OxmlElement("w:bottom")
    b.set(qn("w:val"), "single")
    b.set(qn("w:sz"), "6")
    b.set(qn("w:color"), DIVIDER)
    pBdr.append(b)
    pPr.append(pBdr)
    # Entry 1
    p_ed1_deg = right_doc.add_paragraph()
    r_ed1 = p_ed1_deg.add_run("Bachelor of Business Management – Business Administration")
    run_style(r_ed1, 11, bold=True, color_hex=DARK_BLUE)
    p_ed1_deg.paragraph_format.space_after = Pt(0)
    p_ed1_inst = right_doc.add_paragraph()
    r_ed1_i = p_ed1_inst.add_run("Babson College")
    run_style(r_ed1_i, 10.5, color_hex=BROWN)
    p_ed1_inst.paragraph_format.space_after = Pt(0)
    p_ed1_d = right_doc.add_paragraph()
    r_ed1_d = p_ed1_d.add_run("2002 – 2005")
    run_style(r_ed1_d, 9, color_hex=LIGHT_GRAY)
    p_ed1_d.paragraph_format.space_after = Pt(8)
    # Entry 2
    p_ed2_deg = right_doc.add_paragraph()
    r_ed2 = p_ed2_deg.add_run("Certified Developer – Associate")
    run_style(r_ed2, 11, bold=True, color_hex=DARK_BLUE)
    p_ed2_deg.paragraph_format.space_after = Pt(0)
    p_ed2_inst = right_doc.add_paragraph()
    r_ed2_i = p_ed2_inst.add_run("Georgia Institute of Technology")
    run_style(r_ed2_i, 10.5, color_hex=BROWN)
    p_ed2_inst.paragraph_format.space_after = Pt(0)
    p_ed2_d = right_doc.add_paragraph()
    r_ed2_d = p_ed2_d.add_run("2000 – 2002")
    run_style(r_ed2_d, 9, color_hex=LIGHT_GRAY)

    import os
    out_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "resume_data_engineer.docx")
    doc.save(out_path)
    print("Saved:", out_path)


if __name__ == "__main__":
    main()
